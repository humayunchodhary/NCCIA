#!/usr/bin/env python3
"""Generate NCCIA CMS Infrastructure Specification Word document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date
import os

OUT = os.path.join(os.path.dirname(__file__), '..', 'docs', 'NCCIA-CMS-Infrastructure-Specification-v1.1.docx')


def set_cell_shading(cell, hex_color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.2 * level)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], '1E40AF')
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def build():
    doc = Document()

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = title.add_run('NCCIA Case Management System\n')
    t.bold = True
    t.font.size = Pt(22)
    t.font.color.rgb = RGBColor(30, 64, 175)
    t2 = title.add_run('Complete Infrastructure & Technical Specification\n')
    t2.font.size = Pt(16)
    t2.bold = True
    t3 = title.add_run('OCR at Scale (10 Lakh Files) · SMS Gateway Comparison · National-Level Server Architecture\n')
    t3.font.size = Pt(12)
    t3.italic = True
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m = meta.add_run(f'Prepared for: National Cyber Crime Investigation Agency (NCCIA)\nDocument Date: {date.today().strftime("%d %B %Y")}\nClassification: Official — Internal / Procurement Use\nVersion: 1.0')
    m.font.size = Pt(10)
    doc.add_page_break()

    # 1. Executive Summary
    add_heading(doc, '1. Executive Summary', 1)
    add_para(doc, (
        'This document defines the infrastructure required to operate the NCCIA Case Management System (CMS) '
        'at national scale — including bulk OCR processing of up to 10 lakh (1,000,000) historical and incoming '
        'documents, transactional SMS notifications, and high-availability server architecture for all circles, '
        'HQ, forensic lab integration, and public-facing complaint channels.'
    ))
    add_para(doc, (
        'IMPORTANT — OCR accuracy: No OCR system guarantees 100% accuracy on all document types. Industry practice '
        'is 95–99% field-level accuracy on clean scans, with mandatory human verification for legal/case fields. '
        'This specification targets ≥98% automated extraction with a structured review queue for exceptions.'
    ), bold=False)

    add_heading(doc, '1.1 Current System (As Deployed)', 2)
    add_bullet(doc, 'Web application: Laravel (PHP) API + React SPA')
    add_bullet(doc, 'OCR today: Browser-based Tesseract.js (client-side) for PDF complaint import — suitable for single-user uploads, NOT for 10 lakh batch processing')
    add_bullet(doc, 'SMS today: Generic HTTP gateway (Pakistan telco / SMS aggregator) via SmsService — bilingual EN + Roman Urdu templates')
    add_bullet(doc, 'Database: MySQL/MariaDB with Redis queue recommended for production scale')

    doc.add_page_break()

    # 2. OCR Section
    add_heading(doc, '2. OCR Requirements — 10 Lakh (1,000,000) Files', 1)
    add_heading(doc, '2.1 Scope & Assumptions', 2)
    add_table(doc,
        ['Parameter', 'Assumption', 'Notes'],
        [
            ['Total files', '10,00,000 (10 lakh)', 'PDF scans, complaint forms, FIR copies, annexures'],
            ['Avg pages per file', '3 pages', 'Range 1–12; adjust if mostly single-page'],
            ['Total page OCR volume', '~30,00,000 pages', '10L × 3'],
            ['Target auto-fill accuracy', '≥98% on structured fields', 'CNIC, phone, date, tracking — with human QA queue'],
            ['Processing window', '90–180 days batch + ongoing', 'Parallel worker cluster required'],
            ['Languages', 'English + Urdu (Roman/Nastaliq scans)', 'Urdu Nastaliq needs dedicated models or cloud OCR'],
        ],
        [4, 4, 8]
    )

    add_heading(doc, '2.2 Why Browser OCR Cannot Handle 10 Lakh Files', 2)
    add_bullet(doc, 'Current NCCIA CMS runs Tesseract.js in the user\'s browser — one file at a time, limited RAM, no central queue.')
    add_bullet(doc, '10 lakh files require server-side batch pipeline: upload → queue → OCR workers → validation → CMS auto-fill → audit log.')
    add_bullet(doc, 'Legal admissibility requires hash, timestamp, and version of OCR engine logged per document.')

    add_heading(doc, '2.3 Recommended OCR Architecture (National Scale)', 2)
    add_para(doc, 'Tier A — Production Pipeline (Recommended)', bold=True)
    add_bullet(doc, 'Object storage (MinIO / S3-compatible): store all source PDFs — est. 10L × 2 MB avg = ~2 TB raw')
    add_bullet(doc, 'Message queue: Redis Queue / RabbitMQ / Laravel Horizon — job per page or per file')
    add_bullet(doc, 'OCR Worker nodes: dedicated servers/containers running Tesseract 5.x + optional GPU (CUDA) for speed')
    add_bullet(doc, 'Optional: Azure Document Intelligence / Google Document AI / AWS Textract for Urdu Nastaliq or poor-quality scans')
    add_bullet(doc, 'Post-OCR: regex + ML field extractor → map to CMS complaint/case schema → flag low-confidence rows for manual review')
    add_bullet(doc, 'Admin dashboard: queue depth, pages/hour, error rate, reviewer backlog')

    add_heading(doc, '2.4 OCR Server Sizing — Three Tiers', 2)
    add_table(doc,
        ['Tier', 'Use Case', 'OCR Workers', 'CPU/RAM per worker', 'Est. throughput', 'Batch time (30L pages)'],
        [
            ['Minimum', 'Pilot — 1 lakh files', '4 workers', '8 vCPU, 16 GB RAM', '~8,000 pages/day', '~375 days (too slow)'],
            ['Recommended', 'National — 10 lakh in ~6 months', '16 workers', '16 vCPU, 32 GB RAM', '~55,000 pages/day', '~55 days active + QA'],
            ['Optimal', 'National — 10 lakh in ~90 days', '32 workers + 4× GPU nodes', '16 vCPU + NVIDIA T4/A10', '~120,000 pages/day', '~25 days active + QA'],
        ],
        [2.5, 3, 2, 3, 3, 3.5]
    )

    add_heading(doc, '2.5 OCR Server Specifications (Recommended Tier — Per Worker Node)', 2)
    add_table(doc,
        ['Component', 'Specification'],
        [
            ['CPU', 'Intel Xeon / AMD EPYC — 16 physical cores (32 threads)'],
            ['RAM', '32 GB DDR4 (64 GB if multi-page parallel within node)'],
            ['Storage', '500 GB NVMe SSD (temp render cache) + network mount to object storage'],
            ['GPU (optional)', 'NVIDIA T4 16 GB — 3–5× faster for image preprocessing batches'],
            ['OS', 'Ubuntu 22.04 LTS / RHEL 8'],
            ['Software', 'Tesseract 5.x, Poppler, ImageMagick, Laravel Horizon worker, Docker'],
            ['Network', '1 Gbps minimum to storage and app tier'],
        ],
        [4, 12]
    )

    add_heading(doc, '2.6 OCR Cluster — Total Recommended Hardware (16 Workers)', 2)
    add_table(doc,
        ['Resource', 'Quantity', 'Purpose'],
        [
            ['OCR Worker VM/Bare-metal', '16 nodes', 'Parallel page recognition'],
            ['OCR Master / Queue Server', '1 node (8 vCPU, 16 GB)', 'Redis + Horizon + job dispatch'],
            ['Object Storage Cluster', '3-node MinIO or cloud S3', '~5 TB usable (2 TB docs + growth + OCR cache)'],
            ['Human QA Workstations', '10–20 seats', 'Review low-confidence extractions (<98%)'],
            ['Estimated accuracy after QA', '≥99.5% on committed case data', 'With mandatory sign-off per batch'],
        ],
        [5, 3, 8]
    )

    add_heading(doc, '2.7 OCR Realistic Accuracy Statement', 2)
    add_para(doc, (
        'Claiming "100% OCR accuracy" on 10 lakh mixed-quality government scans is not technically honest. '
        'NCCIA should adopt: (1) Automated extraction ≥98%, (2) Mandatory dual review for P1 cases and legal fields, '
        '(3) Immutable audit log of original PDF hash vs extracted JSON, (4) Officer sign-off before case activation.'
    ))

    doc.add_page_break()

    # 3. SMS Section
    add_heading(doc, '3. SMS Gateway — Package Requirements & Comparison', 1)
    add_heading(doc, '3.1 NCCIA SMS Use Cases', 2)
    add_bullet(doc, 'Complaint acknowledgment with tracking number (bilingual)')
    add_bullet(doc, 'Verification / enquiry / case assignment alerts to officers')
    add_bullet(doc, 'Forensic report ready, court date reminders')
    add_bullet(doc, 'OTP / password reset (if enabled)')
    add_bullet(doc, 'Estimated volume: 50,000–500,000 SMS/month at national scale')

    add_heading(doc, '3.2 Comparison: SMS vs WhatsApp vs AWS SNS', 2)
    add_table(doc,
        ['Criteria', 'Pakistan SMS Gateway\n(Telenor/Zong/UFone API, sms.com.pk, sendpk.com)', 'WhatsApp Business API\n(Meta Cloud API)', 'AWS SNS\n(Simple Notification Service)'],
        [
            ['Reach in Pakistan', 'Excellent — all mobile numbers', 'Good — requires WhatsApp installed + opt-in', 'Limited — needs local SMS origination partner for PK numbers'],
            ['PTA / regulatory', 'Registered alphanumeric sender (e.g. NCCIA)', 'Meta Business verification required', 'AWS account + compliance; often routes via third-party for PK'],
            ['Delivery speed', '5–30 seconds', '5–60 seconds', 'Variable by route'],
            ['Rich content', 'Plain text 160/1600 chars', 'Text, buttons, templates, media', 'Plain text / limited'],
            ['Delivery reports', 'Yes (DLR via gateway)', 'Yes (read receipts optional)', 'Yes (SNS delivery status)'],
            ['Two-way chat', 'Limited (short codes)', 'Full conversational', 'No'],
            ['Official government fit', 'HIGH — standard for citizen alerts', 'MEDIUM — good for engagement, not all citizens', 'MEDIUM — better for cloud-native global apps'],
            ['NCCIA CMS integration', 'ALREADY BUILT — SmsService HTTP gateway', 'Requires new WhatsApp module + templates', 'Requires SNS SDK + PK SMS origination'],
            ['Urdu / Roman Urdu', 'Full Unicode SMS (UCS-2)', 'Full Unicode', 'Full Unicode'],
        ],
        [3.5, 4, 4, 4]
    )

    add_heading(doc, '3.3 Recommended SMS Strategy for NCCIA', 2)
    add_para(doc, 'Primary channel: Pakistan PTA-registered SMS gateway (transactional route)', bold=True)
    add_bullet(doc, 'Use existing SmsService with approved sender ID "NCCIA"')
    add_bullet(doc, 'Contract: minimum 100K–500K SMS/month bundle with telco or aggregator (Telenor Corporate SMS, Zong B2B, Ufone, or licensed aggregator)')
    add_bullet(doc, 'Secondary (optional): WhatsApp Business API for rich notifications where complainant opts in')
    add_bullet(doc, 'AWS SNS: only if entire stack moves to AWS and local SMS origination partner is contracted')

    add_heading(doc, '3.4 SMS Package / Contract Requirements', 2)
    add_table(doc,
        ['Requirement', 'Detail'],
        [
            ['Sender ID', 'Alphanumeric "NCCIA" — PTA approved'],
            ['Route type', 'Transactional (not promotional)'],
            ['API type', 'HTTP GET/POST REST — matches current CMS SmsService'],
            ['Delivery report (DLR)', 'Webhook or poll URL for failed numbers'],
            ['Unicode', 'Required for Roman Urdu'],
            ['Monthly bundle', '250,000 SMS minimum recommended for national ops'],
            ['Failover', 'Secondary gateway URL in config if primary down'],
            ['SLA', '99% delivery within 60 seconds for on-net numbers'],
        ],
        [4, 12]
    )

    add_heading(doc, '3.5 SMS Message Amount — Monthly Estimate (PKR Only)', 2)
    add_para(doc, (
        'Note: Document mein sirf SMS / message ki estimated amount di gayi hai. '
        'Server, OCR, cloud ya kisi aur infrastructure ki cost is section mein shamil nahi.'
    ))
    add_table(doc,
        ['Monthly SMS volume', 'Rate per SMS (PKR)', 'Total message amount (PKR/month)'],
        [
            ['50,000 messages', '1.50', '75,000'],
            ['100,000 messages', '1.40', '140,000'],
            ['250,000 messages', '1.20 (enterprise bundle)', '300,000'],
            ['500,000 messages', '1.00 (national bundle)', '500,000'],
        ],
        [5, 4, 4]
    )
    add_para(doc, (
        'Ye amounts Pakistan transactional SMS gateway (Telenor / Zong / Ufone / licensed aggregator) '
        'par based hain. Actual rate telco contract par depend karega. '
        'Har complaint acknowledgment bilingual (EN + UR) = 2 SMS per event jahan dono bheje jayein.'
    ))

    doc.add_page_break()

    # 4. National Server Architecture
    add_heading(doc, '4. Complete Server Specifications — National Level CMS', 1)
    add_heading(doc, '4.1 Scale Assumptions (National NCCIA)', 2)
    add_table(doc,
        ['Parameter', 'Estimate'],
        [
            ['Concurrent users (peak)', '800 – 1,500 officers'],
            ['Registered users', '3,000 – 8,000'],
            ['Circles / zones', '15 – 25'],
            ['New complaints / month', '5,000 – 25,000'],
            ['Active cases / enquiries', '50,000 – 200,000 live records'],
            ['Document storage (5 years)', '10 – 50 TB (PDF, images, forensic exports)'],
            ['Database size (5 years)', '500 GB – 2 TB MySQL'],
            ['Uptime target', '99.9% (national critical system)'],
        ],
        [6, 6]
    )

    add_heading(doc, '4.2 Production Architecture — High Availability', 2)
    add_para(doc, 'Recommended topology: Active-Passive or Active-Active across two data centres (Islamabad primary + disaster recovery site).')
    add_bullet(doc, 'Load Balancer (×2 HA): HAProxy / F5 / cloud LB — SSL termination, rate limiting')
    add_bullet(doc, 'Web/App tier (×3 minimum): Laravel + PHP 8.2+, Nginx, PHP-FPM')
    add_bullet(doc, 'React static assets: CDN or Nginx on app nodes')
    add_bullet(doc, 'Database: MySQL 8 / MariaDB 10.11 — Primary + synchronous replica + async DR replica')
    add_bullet(doc, 'Redis: Sentinel cluster (3 nodes) — cache, sessions, queues')
    add_bullet(doc, 'Object storage: MinIO cluster or cloud — all uploads, PDFs, forensic files')
    add_bullet(doc, 'OCR worker tier: separate from web tier (Section 2)')
    add_bullet(doc, 'Backup: daily DB + hourly incremental files; off-site encrypted backup')

    add_heading(doc, '4.3 Application Server Specifications (Per Node × 3)', 2)
    add_table(doc,
        ['Component', 'Specification'],
        [
            ['Role', 'Web + API (Laravel) + Queue consumer (light)'],
            ['CPU', '16 vCPU (Intel Xeon Silver or equivalent)'],
            ['RAM', '32 GB'],
            ['Storage', '200 GB NVMe SSD (OS + app) — files on object storage'],
            ['OS', 'Ubuntu 22.04 LTS'],
            ['Web stack', 'Nginx 1.24+, PHP 8.2-FPM, OPcache, Redis extension'],
            ['Quantity', '3 nodes minimum (horizontal scale)'],
        ],
        [4, 12]
    )

    add_heading(doc, '4.4 Database Server Specifications', 2)
    add_table(doc,
        ['Component', 'Primary DB', 'Replica DB', 'DR Site'],
        [
            ['CPU', '32 vCPU', '32 vCPU', '16 vCPU'],
            ['RAM', '128 GB', '128 GB', '64 GB'],
            ['Storage', '2 TB NVMe SSD RAID10', '2 TB NVMe SSD', '2 TB SSD'],
            ['Engine', 'MySQL 8.0 / MariaDB 10.11', 'Read replica', 'Async replica'],
            ['Backup', 'Daily full + binlog', '—', 'Weekly restore test'],
        ],
        [3, 4, 4, 4]
    )

    add_heading(doc, '4.5 Redis / Queue Server', 2)
    add_table(doc,
        ['Component', 'Specification'],
        [
            ['Deployment', '3-node Redis Sentinel HA'],
            ['CPU / RAM per node', '8 vCPU, 16 GB RAM'],
            ['Storage', '100 GB SSD (AOF persistence)'],
            ['Use', 'Session cache, Laravel Horizon queues, rate limits'],
        ],
        [4, 12]
    )

    add_heading(doc, '4.6 File / Object Storage', 2)
    add_table(doc,
        ['Component', 'Specification'],
        [
            ['Type', 'MinIO 4-node erasure coding OR enterprise NAS'],
            ['Usable capacity', '20 TB initial (expandable to 100 TB)'],
            ['Network', '10 Gbps storage network'],
            ['Security', 'Encryption at rest, bucket policies, virus scan on upload'],
        ],
        [4, 12]
    )

    add_heading(doc, '4.7 Network & Security', 2)
    add_table(doc,
        ['Item', 'Requirement'],
        [
            ['Internet bandwidth', '1 Gbps symmetric (primary), 500 Mbps (DR)'],
            ['Firewall', 'Next-gen firewall — WAF for OWASP Top 10'],
            ['SSL', 'Government CA or DigiCert wildcard for *.nccia.gov.pk'],
            ['VPN', 'Site-to-site VPN for circle offices without public CMS access'],
            ['DDoS', 'Cloudflare / ISP scrubbing for public complaint portal'],
            ['SIEM / logs', 'Centralised logging — 90-day retention minimum'],
            ['Penetration test', 'Annual third-party security audit'],
        ],
        [4, 12]
    )

    add_heading(doc, '4.8 Complete Hardware Bill of Materials (Summary)', 2)
    add_table(doc,
        ['#', 'Component', 'Qty', 'Purpose'],
        [
            ['1', 'Load Balancer (HA pair)', '2', 'Traffic distribution, SSL'],
            ['2', 'App / Web Server', '3', 'Laravel + React CMS'],
            ['3', 'Database Primary', '1', 'MySQL master'],
            ['4', 'Database Replica', '1', 'Read scaling + failover'],
            ['5', 'DR Database', '1', 'Disaster recovery site'],
            ['6', 'Redis Sentinel nodes', '3', 'Cache + queues'],
            ['7', 'Object Storage cluster', '4', 'Files, PDFs, exports'],
            ['8', 'OCR Worker nodes', '16', '10 lakh file processing'],
            ['9', 'OCR Queue master', '1', 'Job orchestration'],
            ['10', 'Backup server / tape', '1', 'Off-site backups'],
            ['11', 'Monitoring (Prometheus/Grafana)', '1', 'Uptime + alerts'],
        ],
        [1, 5, 1.5, 7]
    )

    add_heading(doc, '4.9 Virtualisation / Cloud Alternative', 2)
    add_para(doc, 'If hosted on cloud (AWS / Azure / local Pakistani cloud e.g. Nayatel, PTCL Cloud):')
    add_bullet(doc, 'App tier: 3× Standard_D8s_v5 (8 vCPU, 32 GB) or equivalent')
    add_bullet(doc, 'DB: Azure Database for MySQL Flexible Server — 32 vCPU, 128 GB, HA enabled')
    add_bullet(doc, 'Storage: Azure Blob / S3 — 20 TB with lifecycle policies')
    add_bullet(doc, 'OCR: 16× compute-optimised VMs + optional GPU NC-series for OCR burst')
    add_bullet(doc, 'OCR: 16× compute-optimised VMs + optional GPU NC-series for OCR burst')

    doc.add_page_break()

    # 5. Cursor IDE Importance
    add_heading(doc, '5. Cursor IDE — Importance for NCCIA CMS Development & Maintenance', 1)
    add_para(doc, (
        'NCCIA Case Management System Cursor AI-powered IDE ke zariye develop aur maintain ho raha hai. '
        'Cursor ek advanced code editor hai jo AI assistance ke sath Laravel, React, aur database '
        'kaam tez aur secure banata hai — khas tor par national-level government project ke liye.'
    ))

    add_heading(doc, '5.1 Cursor Kya Hai?', 2)
    add_para(doc, (
        'Cursor (cursor.com) Visual Studio Code par based AI IDE hai. Yeh codebase samajh kar '
        'features likhta, bugs fix karta, documentation banata, aur purane code ko safely update '
        'karta hai — human developer ki supervision ke sath.'
    ))

    add_heading(doc, '5.2 NCCIA CMS Mein Cursor Ki Importance', 2)
    add_table(doc,
        ['Area', 'Cursor se kya faida hua / hoga', 'NCCIA module example'],
        [
            ['Tez development', 'Features weeks se days mein — boilerplate, API, React pages auto-generate', 'DSR/D.O. Letter, Case Process History, Reference Library'],
            ['Bug fixing', 'Production 500 errors, route issues, permission bugs jaldi trace', 'DSR create fail, login history blank, Moharrar role fix'],
            ['Security & QA', 'Code review, password fields, permission guards check', 'Role-based sidebar, API middleware, audit logs'],
            ['Documentation', 'Word specs, SOPs, user manuals, infrastructure docs', 'Is document ka generation, Reference Library seed content'],
            ['Consistency', 'Poora codebase ek style mein — Laravel + React conventions match', 'SmsService, ReferenceController, permissions.js'],
            ['Maintenance', 'Server deploy commands, git workflow, migration checks', 'reference:ensure, artisan cache, production pull guide'],
            ['Future scale', 'OCR pipeline, SMS gateway, national server modules plan karna', '10 lakh OCR architecture design in this document'],
        ],
        [3, 6, 5]
    )

    add_heading(doc, '5.3 Kyun Cursor Zaroori Hai (Government Project Ke Liye)', 2)
    add_bullet(doc, 'National CMS jaisa bara system ek choti team se maintain ho sakta hai — Cursor AI developer productivity 3–5× badhata hai.')
    add_bullet(doc, 'Har circle ka alag requirement (DSR, forensic, court, SMS) jaldi implement hota hai bina purane code ko todhe.')
    add_bullet(doc, 'Urdu/Roman Urdu content, bilingual SMS templates, official print layouts — sab ek codebase mein track rehta hai.')
    add_bullet(doc, 'New officers join karein to Cursor se onboarding fast — codebase explain, flow samjhana, safe changes.')
    add_bullet(doc, 'Audit trail: git history + AI-assisted changes reviewable — government accountability ke liye important.')
    add_bullet(doc, 'Long-term: OCR cluster, WhatsApp Phase 2, mobile app — future modules Cursor se plan aur build ho sakte hain.')

    add_heading(doc, '5.4 Recommendation — Cursor Team License', 2)
    add_para(doc, (
        'NCCIA IT team ke liye Cursor Pro / Business licenses recommend hain jo NCCIA CMS maintain '
        'karte hain. Yeh server cost nahi — development tool hai. Iske baghair same features '
        'banane mein zyada developers aur zyada time lagega.'
    ))
    add_table(doc,
        ['Item', 'Detail'],
        [
            ['Tool', 'Cursor IDE (AI-powered)'],
            ['Use', 'NCCIA CMS development, bug fixes, documentation, infrastructure planning'],
            ['Who should use', 'Lead developer, backend (Laravel), frontend (React), DevOps'],
            ['Benefit', 'Faster delivery, fewer production bugs, better documentation for national rollout'],
            ['Note', 'Cursor server/SMS amount nahi — sirf software development productivity tool hai'],
        ],
        [4, 12]
    )

    doc.add_page_break()

    # 6. Deployment checklist
    add_heading(doc, '6. Environment Variables & Integration Checklist', 1)
    add_heading(doc, '6.1 SMS (.env)', 2)
    add_table(doc,
        ['Variable', 'Example / Note'],
        [
            ['SMS_ENABLED', 'true'],
            ['SMS_API_URL', 'https://api.provider.com/sms'],
            ['SMS_HTTP_METHOD', 'get or post'],
            ['SMS_SENDER_ID', 'NCCIA (PTA registered)'],
            ['SMS_USERNAME / SMS_API_KEY', 'From telco contract'],
            ['SMS_SUCCESS_MATCH', 'Success substring from gateway response'],
        ],
        [4, 12]
    )

    add_heading(doc, '6.2 OCR (Future Server Pipeline)', 2)
    add_bullet(doc, 'OCR_ENGINE=tesseract|azure|textract')
    add_bullet(doc, 'OCR_QUEUE=redis — Laravel Horizon workers')
    add_bullet(doc, 'OCR_CONFIDENCE_THRESHOLD=0.85 — below this → manual review queue')
    add_bullet(doc, 'STORAGE_DISK=s3 — object storage for source PDFs')

    add_heading(doc, '6.3 Server Deploy Commands (Current CMS)', 2)
    add_para(doc, 'git pull → composer install → php artisan migrate → php artisan reference:ensure → npm build (if needed) → php artisan config:cache → php artisan route:cache → php artisan horizon (queues)')

    doc.add_page_break()

    # 7. Recommendations
    add_heading(doc, '7. Final Recommendations', 1)
    add_table(doc,
        ['Area', 'Recommendation', 'Priority'],
        [
            ['OCR (10 lakh)', 'Build server-side OCR cluster — 16 workers, object storage, QA dashboard. Do NOT rely on browser Tesseract for bulk.', 'CRITICAL'],
            ['OCR accuracy', 'Target ≥98% auto + human QA for legal fields. Never claim 100% without audit.', 'CRITICAL'],
            ['SMS primary', 'Pakistan transactional SMS gateway with PTA sender "NCCIA" — existing SmsService ready.', 'HIGH'],
            ['WhatsApp', 'Optional Phase 2 for opt-in rich messages — not replacement for SMS.', 'MEDIUM'],
            ['AWS SNS', 'Only if full cloud migration; otherwise local gateway is simpler for PK.', 'LOW'],
            ['App servers', '3× 16 vCPU / 32 GB HA behind load balancer.', 'HIGH'],
            ['Database', 'MySQL 128 GB RAM primary + replica + DR.', 'CRITICAL'],
            ['Storage', '20 TB object storage minimum national.', 'HIGH'],
            ['Cursor IDE', 'Maintain team licenses for NCCIA CMS development, fixes, and national rollout documentation.', 'HIGH'],
            ['Security', 'WAF, VPN for circles, annual pentest, SIEM.', 'CRITICAL'],
        ],
        [3, 9, 2]
    )

    add_heading(doc, '8. Document Approval', 2)
    doc.add_paragraph('Prepared by: ___________________________   Date: ______________')
    doc.add_paragraph('Reviewed by (IT): _______________________   Date: ______________')
    doc.add_paragraph('Approved by (DG NCCIA): _________________   Date: ______________')

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.text = f'NCCIA CMS Infrastructure Specification v1.1 — {date.today().year} — Confidential'
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print(f'Created: {OUT}')


if __name__ == '__main__':
    build()
